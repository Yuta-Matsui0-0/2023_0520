import os
import datetime
import docx
import openpyxl

# ファイルのパスを設定
input_folder = "./Input"
output_folder = "./Output"
output_file_prefix = "ResearchSection_"

# セクション名を設定
section_names = ["ABC", "DEF", "GHI", "JKL"]

# ファイル名を取得するための関数
def get_file_name(path):
    return os.path.splitext(os.path.basename(path))[0]

# Excelファイルを作成する
def create_excel_file(output_path, file_names, search_results):
    workbook = openpyxl.Workbook()
    worksheet = workbook.active

    # セクション名を記入
    for i, section_name in enumerate(section_names):
        worksheet.cell(row=1, column=i+2).value = section_name

    # ファイル名とサーチ結果を記入
    for i, file_name in enumerate(file_names):
        worksheet.cell(row=i+2, column=1).value = file_name
        for j, search_result in enumerate(search_results[i]):
            worksheet.cell(row=i+2, column=j+2).value = "v" if search_result else "-"

    # ファイルを保存
    workbook.save(output_path)

# セクションが存在するかを調べる
def search_sections(doc, section_names):
    result = [False] * len(section_names)
    for i, section_name in enumerate(section_names):
        if doc.tables:
            for table in doc.tables:
                if section_name in table.cell(0, 0).text:
                    result[i] = True
                    break
        if doc.paragraphs:
            for paragraph in doc.paragraphs:
                if section_name in paragraph.text:
                    result[i] = True
                    break
    return result

# メインの処理
if __name__ == "__main__":
    # 現在時刻を取得
    current_time = datetime.datetime.now().strftime("%Y%m%d%H%M%S")

    # Inputフォルダ内のファイルを取得
    input_files = [os.path.join(input_folder, f) for f in os.listdir(input_folder) if f.endswith(".docx")]

    # 各ファイルについてセクションを調査
    search_results = []
    file_names = []
    for input_file in input_files:
        doc = docx.Document(input_file)
        search_result = search_sections(doc, section_names)
        file_names.append(get_file_name(input_file))
        search_results.append(search_result)

    # Excelファイルを作成して保存
    output_file_name = output_file_prefix + current_time + ".xlsx"
    output_path = os.path.join(output_folder, output_file_name)
    create_excel_file(output_path, file_names, search_results)
