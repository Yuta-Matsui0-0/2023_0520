import os
import datetime
from docx import Document

def add_tag_after_heading(doc, heading, tag):
    for paragraph in doc.paragraphs:
        if 'Heading' in paragraph.style.name and paragraph.text == heading:
            index = doc.paragraphs.index(paragraph)
            doc.add_paragraph(tag, style=None)._element = paragraph._element.addnext(doc.paragraphs[index]._element)

def process_word_file(file_path, output_folder):
    # Load the Word document
    doc = Document(file_path)
    
    # Add tags after specified headings
    add_tag_after_heading(doc, "ABCD", "<ABCD></>")
    add_tag_after_heading(doc, "EFGH", "<EFGH></>")
    add_tag_after_heading(doc, "IJKL", "<IJKL></>")
    
    # Save the updated Word document
    file_name = os.path.basename(file_path).split(".")[0]
    doc.save(f"{output_folder}/{file_name}_AddTag.docx")

def process_word_files(input_folder, output_folder):
    # Create output folder
    timestamp = datetime.datetime.now().strftime("%Y%m%d-%H%M%S")
    output_folder = f"{output_folder}/{timestamp}_AddTag"
    os.makedirs(output_folder, exist_ok=True)

    # Process all Word files in the input folder
    for file_name in os.listdir(input_folder):
        if file_name.endswith(".docx"):
            process_word_file(f"{input_folder}/{file_name}", output_folder)

input_folder = "path/to/Input"
output_folder = "path/to/Output"

process_word_files(input_folder, output_folder)
