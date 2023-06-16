import os
import docx
from docx import Document
import datetime
import shutil
import pandas as pd

def add_tag_after_heading(doc, heading, tag, filename):
    tag_added = False
    heading_indices = [i for i, para in enumerate(doc.paragraphs) if 'Heading' in para.style.name and para.text == heading]
    for i in heading_indices:
        new_paragraph = doc.add_paragraph(tag)
        doc.element.body.insert(i+1, new_paragraph._element)
        tag_added = True
    if tag_added:
        if filename not in tag_added_files:
            tag_added_files[filename] = []
        tag_added_files[filename].append(heading)

def export_to_excel(tag_added_files):
    writer = pd.ExcelWriter(os.path.join(output_folder, 'Tag_Added_Report.xlsx'), engine='openpyxl')
    for filename in tag_added_files:
        data = {"Heading": headings, "Tag Added": []}
        for heading in headings:
            if filename in tag_added_files and heading in tag_added_files[filename]:
                data["Tag Added"].append("Yes")
            else:
                data["Tag Added"].append("No")
        df = pd.DataFrame(data)
        df.to_excel(writer, sheet_name=filename, index=False)
    writer.save()

input_folder = './Input'
output_folder = './Output'
headings = ['ABCD', 'EFGH', 'IJKL']
tags = ['<ABCD></>', '<EFGH></>', '<IJKL></>']
tag_added_files = {}

current_timestamp = datetime.datetime.now().strftime("%Y%m%d-%H%M%S")
output_folder = os.path.join(output_folder, f"{current_timestamp}_AddTag")
os.mkdir(output_folder)

for filename in os.listdir(input_folder):
    if filename.endswith('.docx'):
        file_path = os.path.join(input_folder, filename)
        doc = Document(file_path)
        for heading, tag in zip(headings, tags):
            add_tag_after_heading(doc, heading, tag, filename)
        output_file_path = os.path.join(output_folder, filename[:-5] + '_AddTag.docx')
        doc.save(output_file_path)

export_to_excel(tag_added_files)